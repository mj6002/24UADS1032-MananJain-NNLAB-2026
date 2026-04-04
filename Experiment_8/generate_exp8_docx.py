import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_exp8_docx():
    doc = docx.Document()
    
    # Title
    title = doc.add_paragraph()
    run = title.add_run("Experiment 8: Mini Project 2 - Bike Driving Behavior Analysis")
    run.bold = True
    run.font.size = Pt(16)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Author
    author = doc.add_paragraph()
    run = author.add_run("B.E. (AI & DS) VI Semester Student")
    run.italic = True
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 1. Objective
    doc.add_heading('1. Objective', level=1)
    p = doc.add_paragraph(
        "The objective of this project is to analyze motion sensor data (Accelerometer and Gyroscope) "
        "captured during bike riding to evaluate driving behavior. Key goals include building a risk "
        "scoring model and comparing RNN (LSTM/GRU) and Transformer architectures."
    )
    
    # 2. Submission Timeline
    doc.add_heading('2. Submission Timeline', level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Milestone'
    hdr_cells[1].text = 'Submission Date'
    
    milestones = [
        ("Dataset Submission", "26.02.2026"),
        ("Model Submission", "21.03.2026"),
        ("Bonus Submission", "30.03.2026"),
        ("Poster Submission", "04.04.2026")
    ]
    
    for mil, dt in milestones:
        row_cells = table.add_row().cells
        row_cells[0].text = mil
        row_cells[1].text = dt

    # 3. Methodology
    doc.add_heading('3. Methodology', level=1)
    doc.add_paragraph("3.1 Dataset Preparation: Synthetic generation of multi-axis sensor data including variations in rider behavior and phone placement.")
    doc.add_paragraph("3.2 Model Architectures: Implementation of sequential RNN (LSTM) and attention-based Transformer encoders.")
    doc.add_paragraph("3.3 Scoring Logic: Mapping risk predictions (0.0 - 1.0) to a Driving Score (0 - 100).")

    # 4. Results & Performance Comparison (Bonus)
    doc.add_heading('4. Results & Performance Comparison (Bonus)', level=1)
    p = doc.add_paragraph("Comparison of LSTM and Transformer performance on the task:")
    
    res_table = doc.add_table(rows=1, cols=3)
    res_table.style = 'Table Grid'
    hdr = res_table.rows[0].cells
    hdr[0].text = 'Model'
    hdr[1].text = 'Final Val Loss (MSE)'
    hdr[2].text = 'Training Time (s)'
    
    res_data = [
        ("LSTM", "0.0085", "8.54"),
        ("Transformer", "0.0095", "26.18")
    ]
    
    for m, l, t in res_data:
        row = res_table.add_row().cells
        row[0].text = m
        row[1].text = l
        row[2].text = t

    # Add Image
    img_path = os.path.join(os.path.dirname(__file__), "model_comparison_results.png")
    if os.path.exists(img_path):
        doc.add_paragraph("\n[Visualization: Validation Loss and Training Time Comparison]")
        doc.add_picture(img_path, width=Inches(6.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 5. Driving Score Evaluation
    doc.add_heading('5. Driving Score Evaluation', level=1)
    doc.add_paragraph("Example model outputs on test scenarios:")
    doc.add_paragraph("- Safe Trip Recording: 85.53 / 100 (Excellent)")
    doc.add_paragraph("- Risky Trip Recording: 21.90 / 100 (Dangerous)")

    # Save
    output_docx = os.path.join(os.path.dirname(__file__), "Lab_Record_Exp8.docx")
    doc.save(output_docx)
    print(f"Successfully generated {output_docx}")

if __name__ == "__main__":
    create_exp8_docx()
